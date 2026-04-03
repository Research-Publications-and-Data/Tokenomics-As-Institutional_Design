"""
Build B2_Final_v18.docx and B3_GeoDePIN_Final.docx
- Fixes Fig 4 p-value annotation (0.016 → 0.031 two-sided)
- Replaces B2 Figs 4-8 images; inserts Fig 3 (was missing)
- Replaces B3 Figs 1-2 images (Who Burns + GEODNET trajectory)
- Adds GEODNET trajectory narrative sentence to B3 §5.3
- Adds S2R limitations caveat to B3 §6 limitations
Run from ~/b2-governance-data/
"""

import shutil, zipfile, re, io
from pathlib import Path
from copy import deepcopy
from lxml import etree
from PIL import Image
import pandas as pd, numpy as np
from scipy import stats
import matplotlib.pyplot as plt
import matplotlib as mpl

# ── 0. Paths ──────────────────────────────────────────────────────────────────
B2_SRC  = Path('paper/B2_Final_v17.docx')
B2_OUT  = Path('paper/B2_Final_v18.docx')
B3_SRC  = Path('b3/paper/B3_GeoDePIN_PostSurgery.docx')
B3_OUT  = Path('b3/paper/B3_GeoDePIN_Final.docx')
EXH     = Path('exhibits/updated')

NSMAP = {
    'w':  'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
    'a':  'http://schemas.openxmlformats.org/drawingml/2006/main',
    'r':  'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    'pic':'http://schemas.openxmlformats.org/drawingml/2006/picture',
}

# ── 1. Fix Fig 4: two-sided Mann-Whitney p-value ──────────────────────────────
print('Step 1: Regenerating Fig 4 with two-sided p-value...')

plt.rcParams.update({
    'font.family': 'serif',
    'font.serif': ['Times New Roman', 'DejaVu Serif'],
    'font.size': 11, 'axes.titlesize': 13, 'axes.labelsize': 12,
    'xtick.labelsize': 10, 'ytick.labelsize': 10, 'legend.fontsize': 10,
    'figure.dpi': 300, 'savefig.dpi': 300, 'savefig.bbox': 'tight',
    'axes.spines.top': False, 'axes.spines.right': False,
})
DEFI_COLOR, DEPIN_COLOR = '#cb6d51', '#2171b5'

reg   = pd.read_csv('data/processed/regression_data_april2026.csv')
defi  = reg[reg['category'] == 'DeFi']['hhi']
depin = reg[reg['category'] == 'DePIN']['hhi']

_, p_one = stats.mannwhitneyu(depin, defi, alternative='greater')
_, p_two = stats.mannwhitneyu(depin, defi, alternative='two-sided')
d_cohen  = (depin.mean() - defi.mean()) / np.sqrt((depin.std()**2 + defi.std()**2) / 2)

print(f'  one-sided p={p_one:.4f}  two-sided p={p_two:.4f}  d={d_cohen:.3f}')
assert abs(p_two - 0.031) < 0.005, f'Two-sided p={p_two:.4f} (expected ~0.031)'

fig, ax = plt.subplots(figsize=(7, 8))
bp = ax.boxplot([defi.values, depin.values], positions=[1, 2], widths=0.5,
                patch_artist=True, showmeans=True,
                meanprops=dict(marker='D', markerfacecolor='white',
                               markeredgecolor='black', markersize=8),
                medianprops=dict(color='black', linewidth=2))
bp['boxes'][0].set_facecolor(DEFI_COLOR);  bp['boxes'][0].set_alpha(0.4)
bp['boxes'][1].set_facecolor(DEPIN_COLOR); bp['boxes'][1].set_alpha(0.4)
np.random.seed(42)
for data, pos, color in [(defi, 1, DEFI_COLOR), (depin, 2, DEPIN_COLOR)]:
    j = np.random.normal(0, 0.04, len(data))
    ax.scatter(pos + j, data.values, color=color, alpha=0.6, s=40,
               zorder=3, edgecolors='white', linewidth=0.5)
ax.text(1.35, defi.mean(),  f'Mean: {defi.mean():.3f}',  va='center', fontsize=10)
ax.text(2.35, depin.mean(), f'Mean: {depin.mean():.3f}', va='center', fontsize=10)
y_top = max(depin.max(), defi.max()) + 0.01
ax.plot([1, 1, 2, 2], [y_top, y_top+0.005, y_top+0.005, y_top], 'k-', linewidth=1)
ax.text(1.5, y_top+0.008, f'Mann-Whitney p = {p_two:.3f}', ha='center',
        fontsize=10, style='italic')
ax.text(1.5, y_top+0.002, f"Cohen's d = {d_cohen:.2f}", ha='center',
        fontsize=10, style='italic')
ax.set_xticks([1, 2])
ax.set_xticklabels([f'DeFi (N={len(defi)})', f'DePIN (N={len(depin)})'], fontsize=11)
ax.set_ylabel('Governance HHI (post-exclusion)')
fig.text(0.5, -0.02, 'Source: Author calculations. Diamond = mean. '
         'Horizontal line = median. Individual protocols overlaid.',
         ha='center', fontsize=8, style='italic', color='gray')
for ext in ['png', 'pdf']:
    fig.savefig(EXH / f'fig4_sector_boxplot.{ext}', dpi=300, bbox_inches='tight')
plt.close()
print(f'  ✓ fig4_sector_boxplot regenerated (p = {p_two:.3f} two-sided)')


# ── Helpers ───────────────────────────────────────────────────────────────────

def inches_to_emu(inches):
    return int(inches * 914400)

def image_emu(img_path, target_width_emu):
    """Compute height EMU preserving aspect ratio given target width."""
    img = Image.open(img_path)
    w, h = img.size
    return int(target_width_emu * h / w)

def resize_image_to_px(src_path, target_w, target_h):
    """Return PNG bytes resized to target_w × target_h pixels."""
    img = Image.open(src_path).convert('RGBA')
    resized = img.resize((target_w, target_h), Image.LANCZOS)
    buf = io.BytesIO()
    resized.save(buf, format='PNG')
    return buf.getvalue()

def replace_image_in_zip(src_zip, dst_zip, media_target, new_image_bytes):
    """Copy src_zip to dst_zip, replacing media_target with new_image_bytes."""
    with zipfile.ZipFile(src_zip, 'r') as zin, \
         zipfile.ZipFile(dst_zip, 'w', zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            if item.filename == media_target:
                zout.writestr(item, new_image_bytes)
            else:
                zout.writestr(item, zin.read(item.filename))

def multi_replace_and_edit(src_zip, dst_zip, replacements, xml_edits=None):
    """
    replacements: dict {archive_path: new_bytes}
    xml_edits:    list of (archive_path, transform_fn) where transform_fn(xml_str)->str
    """
    with zipfile.ZipFile(src_zip, 'r') as zin, \
         zipfile.ZipFile(dst_zip, 'w', zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename in replacements:
                data = replacements[item.filename]
            if xml_edits:
                for path, fn in xml_edits:
                    if item.filename == path:
                        data = fn(data.decode('utf-8')).encode('utf-8')
            zout.writestr(item, data)


# ── 2. Build B2_Final_v18.docx ────────────────────────────────────────────────
print('\nStep 2: Building B2_Final_v18.docx...')

# Image replacement map: archive path → (new exhibit path, original pixel dims)
# Original pixel dims used for resize to match document layout expectations
B2_IMG_REPLACEMENTS = {
    'word/media/exhibit_sector_comparison.png': (
        EXH / 'fig4_sector_boxplot.png', (1810, 1504)),
    'word/media/exhibit_initial_allocation.png': (
        EXH / 'fig5_allocation_scatter.png', (3000, 1800)),
    'word/media/exhibit_delegation_adjusted.png': (
        EXH / 'fig6_delegation_grouped.png', (3000, 1500)),
    'word/media/exhibit_subsidy_vs_hhi.png': (
        EXH / 'fig7_subsidy_scatter.png', (2520, 1654)),
    'word/media/exhibit_concentration_participation.png': (
        EXH / 'fig8_participation.png', (3000, 1800)),
}

# Also add Fig 3 (HHI bar chart) — it's missing from document body.
# image1.png (rId9) is in the relationships but not embedded.
# We'll replace image1.png and then inject the drawing XML into the document.
B2_IMG_REPLACEMENTS['word/media/image1.png'] = (
    EXH / 'fig3_hhi_bar_40protocols.png', (2969, 1770))

# Build replacement bytes dict
replacements_b2 = {}
for archive_path, (new_img, (orig_w, orig_h)) in B2_IMG_REPLACEMENTS.items():
    replacements_b2[archive_path] = resize_image_to_px(new_img, orig_w, orig_h)
    print(f'  prepared {archive_path.split("/")[-1]} ← {new_img.name}')

# XML transformation for B2 document.xml:
# 1. Update EMU extents for replaced images to preserve aspect ratios
# 2. Insert Fig 3 drawing paragraph before the Figure 3 caption paragraph

def update_extents_for_rid(xml_str, rid_to_cx_cy):
    """
    Update wp:extent and a:ext dimensions for each drawing block containing
    one of the target rIds.
    rid_to_cx_cy: dict {rId: (new_cx, new_cy)}
    Strategy: split on </wp:inline>, find which block contains target rId,
    replace both wp:extent and a:ext cx/cy within that block.
    """
    # Split around each inline drawing block
    blocks = xml_str.split('</wp:inline>')
    for i, block in enumerate(blocks[:-1]):  # last block has no matching inline
        for rid, (new_cx, new_cy) in rid_to_cx_cy.items():
            if f'r:embed="{rid}"' in block:
                # Update wp:extent (appears near start of inline block)
                block = re.sub(
                    r'(<wp:extent\s+cx=")(\d+)("\s+cy=")(\d+)(")',
                    rf'\g<1>{new_cx}\g<3>{new_cy}\g<5>',
                    block
                )
                # Update a:ext (appears in pic:spPr)
                block = re.sub(
                    r'(<a:ext\s+cx=")(\d+)("\s+cy=")(\d+)(")',
                    rf'\g<1>{new_cx}\g<3>{new_cy}\g<5>',
                    block
                )
                blocks[i] = block
                print(f'    ✓ extents updated for {rid}: {new_cx}×{new_cy} EMU')
    return '</wp:inline>'.join(blocks)


def transform_b2_xml(xml_str):
    # --- 1. Update extents for replaced images ---
    rid_to_new_img = {
        'rId16': EXH / 'fig4_sector_boxplot.png',
        'rId12': EXH / 'fig5_allocation_scatter.png',
        'rId13': EXH / 'fig6_delegation_grouped.png',
        'rId17': EXH / 'fig7_subsidy_scatter.png',
        'rId15': EXH / 'fig8_participation.png',
    }
    rid_original_cx = {
        'rId16': 5943600, 'rId12': 5029200, 'rId13': 5029200,
        'rId17': 5943600, 'rId15': 5029200,
    }

    rid_to_cx_cy = {}
    for rid, new_img_path in rid_to_new_img.items():
        orig_cx = rid_original_cx[rid]
        img = Image.open(new_img_path)
        w, h = img.size
        rid_to_cx_cy[rid] = (orig_cx, int(orig_cx * h / w))

    xml_str = update_extents_for_rid(xml_str, rid_to_cx_cy)

    # --- 2. Insert Fig 3 drawing paragraph before Figure 3 caption ---
    # The Figure 3 caption paragraph contains:
    # "Figure 3. Holding-based Herfindahl-Hirschman Index"
    # We need to insert a drawing paragraph BEFORE it.
    # rId9 → image1.png is already in the relationships.
    # New image: 2969×1770 pixels.  Width = 5.5" = 5029200 EMU.
    fig3_img = Image.open(EXH / 'fig3_hhi_bar_40protocols.png')
    fig3_w, fig3_h = fig3_img.size
    fig3_cx = 5029200
    fig3_cy = int(fig3_cx * fig3_h / fig3_w)

    drawing_xml = (
        f'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'<w:pPr><w:jc w:val="center"/></w:pPr>'
        f'<w:r><w:rPr/>'
        f'<w:drawing>'
        f'<wp:inline xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"'
        f' distT="0" distB="0" distL="0" distR="0">'
        f'<wp:extent cx="{fig3_cx}" cy="{fig3_cy}"/>'
        f'<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        f'<wp:docPr id="99" name="fig3_hhi_bar"/>'
        f'<wp:cNvGraphicFramePr>'
        f'<a:graphicFrameLocks xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"'
        f' noChangeAspect="1"/>'
        f'</wp:cNvGraphicFramePr>'
        f'<a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">'
        f'<a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        f'<pic:pic xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        f'<pic:nvPicPr>'
        f'<pic:cNvPr id="99" name="fig3_hhi_bar"/>'
        f'<pic:cNvPicPr/>'
        f'</pic:nvPicPr>'
        f'<pic:blipFill>'
        f'<a:blip xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"'
        f' r:embed="rId9"/>'
        f'<a:stretch><a:fillRect/></a:stretch>'
        f'</pic:blipFill>'
        f'<pic:spPr>'
        f'<a:xfrm><a:off x="0" y="0"/><a:ext cx="{fig3_cx}" cy="{fig3_cy}"/></a:xfrm>'
        f'<a:prstGeom prst="rect"><a:avLst/></a:prstGeom>'
        f'</pic:spPr>'
        f'</pic:pic>'
        f'</a:graphicData>'
        f'</a:graphic>'
        f'</wp:inline>'
        f'</w:drawing>'
        f'</w:r>'
        f'</w:p>'
    )

    # Find the Figure 3 caption paragraph and insert drawing before it
    anchor_text = 'Figure 3. Holding-based Herfindahl'
    if anchor_text in xml_str:
        # Find the <w:p> that contains this text
        # Walk backwards from anchor to find the paragraph start
        idx = xml_str.index(anchor_text)
        # Find the <w:p that starts this paragraph (last <w:p before anchor)
        p_start = xml_str.rfind('<w:p ', 0, idx)
        if p_start == -1:
            p_start = xml_str.rfind('<w:p>', 0, idx)
        xml_str = xml_str[:p_start] + drawing_xml + xml_str[p_start:]
        print('  ✓ Fig 3 drawing paragraph inserted before Figure 3 caption')
    else:
        print('  ⚠️  Fig 3 caption anchor not found — image not inserted')

    return xml_str

# Also update image1.png EMU for Fig 3 in rels (no change needed — rId9 already points to image1.png)

multi_replace_and_edit(
    B2_SRC, B2_OUT,
    replacements=replacements_b2,
    xml_edits=[('word/document.xml', transform_b2_xml)]
)
print(f'  ✓ {B2_OUT} written ({B2_OUT.stat().st_size / 1024:.0f} KB)')


# ── 3. Build B3_GeoDePIN_Final.docx ──────────────────────────────────────────
print('\nStep 3: Building B3_GeoDePIN_Final.docx...')

# B3 image mapping:
# word/media/image2.png (rId10) = Fig 1 in doc = Helium Who Burns
# word/media/image3.png (rId11) = Fig 2 in doc = GEODNET trajectory
# word/media/image4.png (rId12) = Fig 3 in doc = 2×2 matrix (no change)

B3_IMG_REPLACEMENTS = {
    'word/media/image2.png': (EXH / 'b3_fig2_who_burns.png',         (2370, 1655)),
    'word/media/image3.png': (EXH / 'b3_fig3_geodnet_trajectory.png', (2544, 1507)),
}

replacements_b3 = {}
for archive_path, (new_img, (orig_w, orig_h)) in B3_IMG_REPLACEMENTS.items():
    replacements_b3[archive_path] = resize_image_to_px(new_img, orig_w, orig_h)
    print(f'  prepared {archive_path.split("/")[-1]} ← {new_img.name}')

# Text to insert after para 203 (GEODNET trajectory narrative)
GEODNET_TRAJECTORY_TEXT = (
    'The trajectory is not monotonic: accelerating vesting unlocks in '
    'Q2 2025 (monthly unlocks rose from approximately 23\u202fmillion to '
    '30\u202fmillion tokens as team and investor allocations vested) '
    'temporarily overwhelmed subscription burn growth, collapsing the '
    'absorption rate to below 20% despite a scheduled June\u202f30 halving '
    'that cut mining rewards by 50%. Subscription growth subsequently '
    'restored absorption to 36% by February\u202f2026. '
    'This pattern illustrates a supply-side risk absent from the health '
    'dashboard\u2019s current specification: vesting-driven dilution '
    'operates independently of emission schedules and can overwhelm '
    'organic demand even when mining rewards are contracting.'
)

# S2R limitations caveat to insert after "Premature Mechanism Complexity" paragraph
S2R_LIMITATIONS_TEXT = (
    'The S2R metric as currently specified measures burns against mining '
    'emissions only; it does not capture vesting-driven dilution from team, '
    'investor, and ecosystem unlocks. For GEODNET, on-chain vesting '
    'represented approximately 54% of total token supply expansion by '
    'Q2\u202f2025, meaning the realized subsidy ratio was substantially '
    'higher than the mining-emissions-only S2R figure suggests. A '
    'comprehensive fiscal sustainability metric would incorporate all '
    'supply sources, not only algorithmic emissions.'
)

def transform_b3_xml(xml_str):
    # Helper: insert a new plain paragraph after a given anchor phrase
    def insert_para_after(xml, anchor_phrase, para_text):
        if anchor_phrase not in xml:
            print(f'  ⚠️  Anchor not found: "{anchor_phrase[:60]}"')
            return xml

        idx = xml.index(anchor_phrase)
        # Find the end of the paragraph containing this anchor
        p_end_idx = xml.index('</w:p>', idx) + len('</w:p>')

        # Harvest run properties from the paragraph for style matching
        para_start = xml.rfind('<w:p ', 0, idx)
        if para_start == -1:
            para_start = xml.rfind('<w:p>', 0, idx)
        para_chunk = xml[para_start:p_end_idx]

        rpr_match = re.search(r'<w:rPr>.*?</w:rPr>', para_chunk, re.DOTALL)
        rpr = rpr_match.group(0) if rpr_match else ''

        ppr_match = re.search(r'<w:pPr>.*?</w:pPr>', para_chunk, re.DOTALL)
        ppr = ppr_match.group(0) if ppr_match else ''

        # Escape XML special chars in text
        safe_text = (para_text
                     .replace('&', '&amp;')
                     .replace('<', '&lt;')
                     .replace('>', '&gt;')
                     .replace('"', '&quot;'))

        new_para = (
            f'<w:p>'
            f'{ppr}'
            f'<w:r>{rpr}'
            f'<w:t xml:space="preserve">{safe_text}</w:t>'
            f'</w:r>'
            f'</w:p>'
        )
        return xml[:p_end_idx] + new_para + xml[p_end_idx:]

    # Insertion 1: GEODNET trajectory text after para 203
    # Anchor: unique phrase from para 203
    anchor_geodnet = 'approaching Hivemapper'
    if anchor_geodnet not in xml_str:
        # Try alternative anchor
        anchor_geodnet = '0.62 by early 2026'
    if anchor_geodnet not in xml_str:
        anchor_geodnet = 'growth-phase tier'
    xml_str = insert_para_after(xml_str, anchor_geodnet, GEODNET_TRAJECTORY_TEXT)
    if anchor_geodnet in xml_str:
        print(f'  ✓ GEODNET trajectory text inserted (anchor: "{anchor_geodnet}")')

    # Insertion 2: S2R limitations caveat after "Premature Mechanism Complexity" para
    anchor_limits = 'Premature Mechanism Complexity'
    xml_str = insert_para_after(xml_str, anchor_limits, S2R_LIMITATIONS_TEXT)
    if anchor_limits in xml_str:
        print(f'  ✓ S2R limitations caveat inserted (anchor: "{anchor_limits}")')

    # Update EMU extents for replaced images to match new aspect ratios
    b3_rid_to_cx_cy = {}
    for rid, new_img_path in {
        'rId10': EXH / 'b3_fig2_who_burns.png',
        'rId11': EXH / 'b3_fig3_geodnet_trajectory.png',
    }.items():
        orig_cx = 4572000
        img = Image.open(new_img_path)
        w, h = img.size
        b3_rid_to_cx_cy[rid] = (orig_cx, int(orig_cx * h / w))

    xml_str = update_extents_for_rid(xml_str, b3_rid_to_cx_cy)
    return xml_str

multi_replace_and_edit(
    B3_SRC, B3_OUT,
    replacements=replacements_b3,
    xml_edits=[('word/document.xml', transform_b3_xml)]
)
print(f'  ✓ {B3_OUT} written ({B3_OUT.stat().st_size / 1024:.0f} KB)')


# ── 4. Verify ─────────────────────────────────────────────────────────────────
print('\nStep 4: Verification...')

from docx import Document

# Verify B2
b2 = Document(str(B2_OUT))
b2_text = ' '.join(p.text for p in b2.paragraphs)
b2_images = sum(1 for r in b2.part.rels.values() if 'image' in r.reltype.lower())

# Verify B3
b3 = Document(str(B3_OUT))
b3_text = ' '.join(p.text for p in b3.paragraphs)
b3_images = sum(1 for r in b3.part.rels.values() if 'image' in r.reltype.lower())

checks = {
    # B2 core stats
    'B2 p=0.031 in text':       'p = 0.031' in b2_text or '0.031' in b2_text,
    'B2 d=0.96 in text':        'd = 0.96' in b2_text or '0.96' in b2_text,
    'B2 DeFi mean 0.043':       '0.043' in b2_text,
    'B2 DePIN mean 0.090':      '0.090' in b2_text,
    'B2 Pfeffer cited':         'Pfeffer' in b2_text,
    'B2 Ballandies cited':      'Ballandies' in b2_text,
    'B2 images ≥ 8':            b2_images >= 8,
    # B3 core content
    'B3 S2R 2.06':              '2.06' in b3_text,
    'B3 Resource Dependence':   'Resource Dependence' in b3_text,
    'B3 title Who Burns':       'Who Burns' in b3_text,
    'B3 images ≥ 4':            b3_images >= 4,
    # New insertions
    'B3 vesting-driven dilution': 'vesting-driven dilution' in b3_text,
    'B3 June halving':           'halving' in b3_text and 'mining rewards' in b3_text,
    'B3 S2R limitations caveat': '54%' in b3_text and 'vesting' in b3_text,
    # Output files exist
    'B2_Final_v18.docx exists': B2_OUT.exists(),
    'B3_GeoDePIN_Final exists':  B3_OUT.exists(),
}

all_pass = True
for name, passed in checks.items():
    status = '✓' if passed else '✗'
    if not passed:
        all_pass = False
    print(f'  {status}  {name}')

print(f'\n  B2: {len(b2_text.split()):,} words, {b2_images} images')
print(f'  B3: {len(b3_text.split()):,} words, {b3_images} images')
print(f'\n{"  ALL CHECKS PASS" if all_pass else "  FAILURES DETECTED — see ✗ above"}')

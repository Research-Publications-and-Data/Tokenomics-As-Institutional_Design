#!/usr/bin/env python3
"""Apply B2 N=20 paper patches to B2_Final_v2.docx using actual regression results."""

from docx import Document
import json
import re
import os

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DOCX_PATH = os.path.join(BASE, "B2_Final_v2.docx")

# Load actual regression results
with open(os.path.join(BASE, "data", "regression_results.json")) as f:
    R = json.load(f)

# Helper to do find-replace across all paragraphs
def find_replace_text(doc, old, new, count=0):
    """Replace text across paragraph runs. Returns number of replacements."""
    replacements = 0
    for para in doc.paragraphs:
        full = para.text
        if old in full:
            # Rebuild runs with replacement
            new_text = full.replace(old, new, 1 if count == 0 else count)
            if new_text != full:
                # Clear all runs and set first run's text
                for i, run in enumerate(para.runs):
                    if i == 0:
                        run.text = new_text
                    else:
                        run.text = ""
                replacements += 1
    # Also check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    full = para.text
                    if old in full:
                        new_text = full.replace(old, new, 1 if count == 0 else count)
                        if new_text != full:
                            for i, run in enumerate(para.runs):
                                if i == 0:
                                    run.text = new_text
                                else:
                                    run.text = ""
                            replacements += 1
    return replacements

def insert_after_paragraph(doc, search_text, new_text):
    """Insert a new paragraph after the one containing search_text."""
    for i, para in enumerate(doc.paragraphs):
        if search_text in para.text:
            # Create new paragraph after this one
            new_para = doc.paragraphs[i]._element
            from docx.oxml.ns import qn
            from copy import deepcopy
            new_p = deepcopy(new_para)
            # Clear and set text
            for child in list(new_p):
                if child.tag.endswith('}r'):
                    new_p.remove(child)
            run = new_p.makeelement(qn('w:r'), {})
            rPr_source = new_para.findall(qn('w:r'))
            if rPr_source:
                rPr = rPr_source[0].find(qn('w:rPr'))
                if rPr is not None:
                    run.append(deepcopy(rPr))
            t = run.makeelement(qn('w:t'), {})
            t.text = new_text
            t.set(qn('xml:space'), 'preserve')
            run.append(t)
            new_p.append(run)
            new_para.addnext(new_p)
            return True
    return False


def main():
    doc = Document(DOCX_PATH)

    # Build replacement values from actual results
    pearson_r = R["cross_pearson"]["r"]
    pearson_p = R["cross_pearson"]["p"]
    spearman_rho = R["cross_spearman"]["rho"]
    spearman_p = R["cross_spearman"]["p"]
    nolp_r = R["excl_livepeer"]["r"]
    nolp_p = R["excl_livepeer"]["p"]
    depin_r = R["depin_pearson"]["r"]
    depin_p = R["depin_pearson"]["p"]
    depin_rho = R["depin_spearman"]["rho"]
    depin_sp = R["depin_spearman"]["p"]
    defi_r = R["defi_pearson"]["r"]
    defi_p = R["defi_pearson"]["p"]
    log_r = R["cross_log"]["r"]
    log_p = R["cross_log"]["p"]
    sector_p = R["sector_mw"]["p"]

    edits = 0

    # Edit 1: Abstract - sample size
    n = find_replace_text(doc, "35 protocols across four categories", "38 protocols across four categories")
    edits += n
    print(f"Edit 1 (abstract 35->38): {n} replacements")

    # Edit 2: Abstract - subsidy finding
    old_sub = "on-chain subsidy ratio vs HHI: r = +0.15, p = 0.57, N = 17"
    new_sub = f"on-chain subsidy ratio vs HHI: Pearson r = +{pearson_r:.2f}, p = {pearson_p:.3f}, N = 20, but Spearman rho = +{spearman_rho:.2f}, p = {spearman_p:.2f}, indicating the linear result is driven by a single extreme observation"
    n = find_replace_text(doc, old_sub, new_sub)
    edits += n
    print(f"Edit 2 (abstract subsidy): {n} replacements")

    # Edit 3: Update remaining "35 protocols" references
    n = find_replace_text(doc, "35 protocols", "38 protocols")
    edits += n
    print(f"Edit 3 (35->38 protocols): {n} additional replacements")

    # Edit 10a: Old Pearson r=+0.15 in body
    old_r15 = "r = +0.15, p = 0.57, N = 17"
    new_r = f"Pearson r = +{pearson_r:.2f}, p = {pearson_p:.3f}, N = 20"
    n = find_replace_text(doc, old_r15, new_r)
    edits += n
    print(f"Edit 10a (r=+0.15->actual): {n} replacements")

    # Also try variant formatting
    for variant in ["r = +0.15, p = 0.57", "r=+0.15, p=0.57", "r = 0.15, p = 0.57"]:
        n = find_replace_text(doc, variant, f"Pearson r = +{pearson_r:.2f}, p = {pearson_p:.3f}")
        if n > 0:
            edits += n
            print(f"  variant '{variant}': {n} replacements")

    # Edit 10b: Old Spearman
    for variant in ["Spearman rho = +0.37, p = 0.14", "Spearman rho = 0.37, p = 0.14", "rho = +0.37, p = 0.14"]:
        n = find_replace_text(doc, variant, f"Spearman rho = +{spearman_rho:.2f}, p = {spearman_p:.2f}")
        if n > 0:
            edits += n
            print(f"Edit 10b (Spearman): {n} replacements")

    # Edit 10c: N=17 subsidy
    n = find_replace_text(doc, "17 protocols with computable ratios", "20 protocols with computable ratios")
    edits += n
    print(f"Edit 10c (17->20 protocols): {n} replacements")

    n = find_replace_text(doc, "N = 17", "N = 20")
    edits += n
    print(f"Edit 10c (N=17->N=20): {n} replacements")

    # Edit 10d: 8 DePIN -> 11 DePIN
    n = find_replace_text(doc, "8 DePIN", "11 DePIN")
    edits += n
    print(f"Edit 10d (8->11 DePIN): {n} replacements")

    # Edit 10e: Sector test p update
    n = find_replace_text(doc, "p = 0.20", f"p = {sector_p:.2f}")
    edits += n
    print(f"Edit 10e (sector p): {n} replacements")

    # Edit 5: Replace subsidy paragraph
    old_para_start = "On-chain subsidy ratios, computed from protocol-native revenue and emission events"
    new_para = (
        f"On-chain subsidy ratios, computed from protocol-native revenue and emission events for "
        f"11 DePIN and 9 DeFi protocols, show a significant linear association with governance "
        f"concentration cross-sector (Pearson r = +{pearson_r:.2f}, p = {pearson_p:.3f}, N = 20). "
        f"However, this result does not survive rank-based testing (Spearman rho = +{spearman_rho:.2f}, "
        f"p = {spearman_p:.2f}), log transformation (r = +{log_r:.2f}, p = {log_p:.2f}), or "
        f"exclusion of the most leveraged observation (Livepeer, 88.5x subsidy; excluding: "
        f"r = {nolp_r:+.2f}, N = 19). Within DePIN, the rank correlation is effectively zero "
        f"(Spearman rho = {depin_rho:+.2f}, p = {depin_sp:.2f}, N = 11). The apparent linear signal "
        f"reflects one extreme observation rather than a systematic relationship across the sample."
    )
    # Find and replace the paragraph
    for para in doc.paragraphs:
        if old_para_start in para.text:
            for i, run in enumerate(para.runs):
                if i == 0:
                    run.text = new_para
                else:
                    run.text = ""
            edits += 1
            print(f"Edit 5 (subsidy paragraph): replaced")
            break

    # Edit 4: Insert Aethir case study after Livepeer TT measurement text
    aethir_text = (
        "Aethir illustrates the inverse limitation: despite generating the highest DePIN revenue "
        "in the sample ($166M annualized from 150+ enterprise AI and gaming clients), its governance "
        "token (ATH) has an HHI of 0.168, the second highest in the DePIN subsample. Aethir's "
        "subsidy ratio of 0.15x (revenue far exceeds emissions) places it at the opposite extreme "
        "from Livepeer (88.5x), yet both exhibit high governance concentration. This decoupling of "
        "financial health from governance decentralization underscores that concentration is driven "
        "by initial allocation and vesting structures, not by subsidy dependence."
    )
    # Try to find a paragraph mentioning Livepeer + Token Terminal to insert after
    for para in doc.paragraphs:
        if "Livepeer" in para.text and ("Token Terminal" in para.text or "measurement" in para.text.lower()):
            inserted = insert_after_paragraph(doc, para.text[:50], aethir_text)
            if inserted:
                edits += 1
                print("Edit 4 (Aethir case study): inserted")
            break

    # Edit 7: Add robustness paragraph after existing robustness text
    robustness_text = (
        f"The subsidy-concentration association is similarly fragile. The cross-sector Pearson "
        f"correlation (r = +{pearson_r:.2f}, p = {pearson_p:.3f}, N = 20) is driven entirely by "
        f"Livepeer (subsidy ratio 88.5x, HHI 0.199). Removing this single observation eliminates "
        f"the association (r = {nolp_r:+.2f}, p = {nolp_p:.2f}, N = 19). The Spearman rank "
        f"correlation, which is robust to extreme values, shows no relationship (rho = +{spearman_rho:.2f}, "
        f"p = {spearman_p:.2f}). Within the 11-protocol DePIN subsample, the rank correlation is "
        f"effectively zero (rho = {depin_rho:+.2f}, p = {depin_sp:.2f}). The linear significance is "
        f"an artifact of one protocol's extreme leverage on the regression line, not evidence of a "
        f"systematic mechanism linking subsidy structure to governance outcomes."
    )
    for para in doc.paragraphs:
        if "alternative concentration metrics" in para.text.lower() or "robustness" in para.text.lower():
            inserted = insert_after_paragraph(doc, para.text[:50], robustness_text)
            if inserted:
                edits += 1
                print("Edit 7 (robustness paragraph): inserted")
            break

    # Edit 8: Update third contribution
    old_contrib = "Third, no measured covariate predicts"
    new_contrib = (
        f"Third, no measured covariate robustly predicts long-run governance concentration: not "
        f"sector membership, not initial token allocation (r = +0.08, p = 0.64, N = 33), not "
        f"protocol maturity, and not subsidy structure. On-chain subsidy ratios, computed from "
        f"protocol-native revenue and emission events for 11 DePIN and 9 DeFi protocols, show a "
        f"nominal linear association (Pearson r = +{pearson_r:.2f}, p = {pearson_p:.3f}, N = 20) "
        f"that does not survive rank-based testing (Spearman rho = +{spearman_rho:.2f}, "
        f"p = {spearman_p:.2f}) or exclusion of a single extreme observation (Livepeer; excluding: "
        f"r = {nolp_r:+.2f}). Aethir, the highest-revenue DePIN protocol ($166M annualized, "
        f"subsidy 0.15x), has the second-highest governance concentration (HHI 0.168), "
        f"demonstrating that financial sustainability does not ensure governance decentralization."
    )
    for para in doc.paragraphs:
        if old_contrib in para.text:
            for i, run in enumerate(para.runs):
                if i == 0:
                    run.text = new_contrib
                else:
                    run.text = ""
            edits += 1
            print("Edit 8 (third contribution): replaced")
            break

    # Edit 12: Add Livepeer outlier disclosure to limitations
    livepeer_disclosure = (
        "The cross-sector subsidy analysis includes one extreme observation (Livepeer, subsidy "
        "ratio 88.5x) that exerts disproportionate leverage on the Pearson correlation. While the "
        "protocol is not excluded from the main specification (its extreme subsidy is a genuine "
        "feature of its tokenomics, not a measurement error), readers should note that the nominal "
        f"Pearson significance (p = {pearson_p:.3f}) is entirely attributable to this single data "
        f"point. The Spearman rank correlation, which downweights extreme values, shows no "
        f"relationship (p = {spearman_p:.2f}). Future work with larger DePIN samples would provide "
        f"a more definitive test."
    )
    for para in doc.paragraphs:
        if "limitation" in para.text.lower() and ("section" in para.text.lower() or len(para.text) < 50):
            inserted = insert_after_paragraph(doc, para.text[:30], livepeer_disclosure)
            if inserted:
                edits += 1
                print("Edit 12 (Livepeer disclosure): inserted")
            break

    # Edit 6: N=22 -> check
    n = find_replace_text(doc, "N = 22", "N = 22")  # Keep as-is per patch instructions
    print(f"Edit 6 (N=22): kept as-is (expansion protocols lack full covariates)")

    # Edit 11: Keywords - try to find and update
    for para in doc.paragraphs:
        if "Keywords" in para.text or "keywords" in para.text:
            if "sample expansion" not in para.text:
                for run in para.runs:
                    if run.text.strip().endswith('.') or run.text.strip().endswith(';'):
                        run.text = run.text.rstrip('.;') + "; sample expansion; outlier robustness."
                        edits += 1
                        print("Edit 11 (keywords): updated")
                        break
            break

    # Save
    doc.save(DOCX_PATH)
    print(f"\nTotal edits applied: {edits}")
    print(f"Saved to {DOCX_PATH}")

    # Verification
    print("\n=== VERIFICATION ===")
    doc2 = Document(DOCX_PATH)
    full_text = "\n".join(p.text for p in doc2.paragraphs)
    checks = {
        "38 protocols": "38 protocols" in full_text,
        "N=20": "N = 20" in full_text,
        "11 DePIN": "11 DePIN" in full_text,
        f"Pearson r=+{pearson_r:.2f}": f"+{pearson_r:.2f}" in full_text,
        f"p={pearson_p:.3f}": f"{pearson_p:.3f}" in full_text,
        "Spearman rho": f"+{spearman_rho:.2f}" in full_text,
        "Aethir": "Aethir" in full_text,
        "No old r=+0.15": "r = +0.15, p = 0.57" not in full_text,
        "No old N=17 subsidy": full_text.count("N = 17") <= 1,
        "88.5x": "88.5x" in full_text or "88.5" in full_text,
        "Hivemapper": "Hivemapper" in full_text,
    }
    for check, passed in checks.items():
        status = "PASS" if passed else "FAIL"
        print(f"  [{status}] {check}")


if __name__ == "__main__":
    main()

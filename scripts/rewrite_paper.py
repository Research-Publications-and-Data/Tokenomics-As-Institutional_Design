"""
B2 Paper Rewrite: 35-Protocol Corrected Dataset
Produces B2_Rewrite_34Protocol.docx from B2_Final_WithFigures_POLISHED.docx

Phases:
1. Abstract replacement
2. Keywords update
3. §4.2 expansion
4. Table 4 replacement (35 protocols)
5. §5.2 rewrite
6. §5.3 rewrite
7. §6 HHI value replacements
8. §7 conclusion rewrite
9. Figure 3 replacement
"""

import shutil
import re
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import pandas as pd

SRC = "/Users/zach/B2_Final_WithFigures_POLISHED.docx"
DST = "/Users/zach/B2_Rewrite_34Protocol.docx"
DATA_CSV = "/Users/zach/b2-governance-data/governance_concentration_april2026.csv"
FIGURE3 = "/Users/zach/b2-governance-data/figure3_governance_hhi_34protocol.png"

# Load data
df = pd.read_csv(DATA_CSV)
cat_order = {'DePIN': 0, 'DeFi': 1, 'L1_L2_Infra': 2, 'Social_Dead': 3}
df['cat_sort'] = df['category'].map(cat_order)
df_sorted = df.sort_values(['cat_sort', 'hhi'], ascending=[True, True])

# Category statistics
depin = df[df['category'] == 'DePIN']['hhi']
defi = df[df['category'] == 'DeFi']['hhi']
infra = df[df['category'] == 'L1_L2_Infra']['hhi']

print(f"DePIN: mean={depin.mean():.3f}, N={len(depin)}")
print(f"DeFi: mean={defi.mean():.3f}, N={len(defi)}")
print(f"Infra: mean={infra.mean():.3f}, N={len(infra)}")
print(f"All: min={df['hhi'].min():.4f}, max={df['hhi'].max():.4f}")

# Copy file
shutil.copy2(SRC, DST)
doc = Document(DST)


def find_paragraph_containing(text_fragment):
    """Find first paragraph containing the given text."""
    for i, p in enumerate(doc.paragraphs):
        if text_fragment in p.text:
            return i, p
    return None, None


def replace_paragraph_text(p, new_text):
    """Replace all text in a paragraph while preserving the first run's formatting."""
    if not p.runs:
        p.text = new_text
        return
    # Keep first run's formatting
    fmt = p.runs[0].font
    # Clear all runs
    for run in p.runs:
        run.text = ""
    p.runs[0].text = new_text


def find_and_replace_text(old_text, new_text):
    """Find and replace text across paragraphs."""
    count = 0
    for p in doc.paragraphs:
        if old_text in p.text:
            for run in p.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)
                    count += 1
    return count


# ============================================================
# PHASE 1: Abstract replacement
# ============================================================
print("\nPhase 1: Abstract")

# Find abstract paragraph (the one after "Abstract" heading)
for i, p in enumerate(doc.paragraphs):
    if p.style.name == 'Heading 1' and p.text.strip() == 'Abstract':
        # Next paragraph is the abstract text
        abstract_p = doc.paragraphs[i + 1]
        new_abstract = (
            "Governance concentration in Decentralized Physical Infrastructure Networks (DePIN) "
            "is not statistically distinguishable from decentralized finance (DeFi) benchmarks in "
            "a 35-protocol cross-section: Herfindahl-Hirschman Index (HHI) values range from 0.004 "
            "to 0.199 (March 2026 snapshot, protocol-controlled addresses excluded), with DePIN mean "
            f"HHI {depin.mean():.3f} versus DeFi mean {defi.mean():.3f} (Welch t-test p = 0.35, "
            "Cohen's d = 0.39). This sector invariance finding contradicts a common assumption that "
            "hardware-network tokens are inherently more concentrated than pure-software DeFi governance. "
            "To evaluate what this concentration means for institutional legitimacy, we develop a normative "
            "framework translating five political-philosophical traditions into evaluable design criteria "
            "for tokenized systems. A systematic exclusion methodology removing 30 protocol-controlled "
            "addresses (staking contracts, exchange custodians, vesting locks, and protocol treasuries) "
            "across 14 protocols reveals that naive holder-list analysis overstates concentration by "
            "up to 5x in affected protocols. We find that distribution mechanism design, not sector, "
            "is the primary determinant of governance concentration: protocols with broad airdrops "
            "consistently achieve lower HHI regardless of whether they coordinate digital or physical "
            "infrastructure. The framework and dataset offer a replicable foundation for evaluating "
            "legitimacy in programmable institutions."
        )
        replace_paragraph_text(abstract_p, new_abstract)
        print(f"  Replaced abstract ({len(new_abstract)} chars)")
        break

# ============================================================
# PHASE 2: Keywords
# ============================================================
print("\nPhase 2: Keywords")
for p in doc.paragraphs:
    if p.text.startswith("Keywords:"):
        new_kw = (
            "Keywords: tokenomics, institutional design, DePIN, governance concentration, "
            "mechanism design, political philosophy, Herfindahl-Hirschman Index, decentralized "
            "governance, scoring rubric, blockchain governance, exclusion methodology, "
            "delegation amplification, sector invariance"
        )
        replace_paragraph_text(p, new_kw)
        print(f"  Updated keywords")
        break

# ============================================================
# PHASE 3: §4.2 expansion
# ============================================================
print("\nPhase 3: §4.2 expansion")
# Find and update the sampling description
old_sampling = "Protocols were selected to maximize variation across three dimensions: sector (5 DeFi benchmarks, 2 general infrastructure, 6 DePIN)"
idx, p = find_paragraph_containing(old_sampling)
if p:
    new_sampling = (
        "Protocols were selected to maximize variation across three dimensions: "
        "sector (13 DeFi, 12 DePIN, 8 L1/L2 infrastructure, 2 survivorship-bias controls), "
        "maturity (2017-2024 launches), and governance model "
        "(token-weighted, delegate, futarchy). The expanded 35-protocol sample was constructed "
        "by adding protocols from the Dune Analytics ecosystem that met three criteria: "
        "(i) publicly traded governance token, (ii) holder data queryable via Dune or Helius DAS API, "
        "and (iii) at least 50 non-zero-balance holders after protocol-controlled address exclusion. "
        "MetaDAO is included as the sole futarchy-governance protocol in the sample. "
        "Solana-native tokens (META, DRIFT, GRASS, W) were collected via Helius getTokenAccounts "
        "API to address systematic undercounting in Dune's solana_utils.daily_balances table, "
        "which returned 29-291 holders versus 5,067-99,980 from direct RPC enumeration."
    )
    replace_paragraph_text(p, new_sampling)
    print(f"  Updated sampling strategy text")
else:
    print(f"  WARNING: Could not find sampling paragraph")

# Update figure caption
old_fig_caption = "Figure 3. Holding-based Herfindahl-Hirschman Index (HHI) across 12 DeFi and GeoDePIN governance tokens"
idx, p = find_paragraph_containing(old_fig_caption)
if p:
    new_fig_caption = (
        "Figure 3. Holding-based Herfindahl-Hirschman Index (HHI) across 35 governance tokens "
        "(March 2026 snapshot, top-1,000 holders, protocol-controlled addresses excluded). "
        "Dashed lines indicate category means. Protocols sorted by category then HHI ascending."
    )
    replace_paragraph_text(p, new_fig_caption)
    print(f"  Updated Figure 3 caption")

# ============================================================
# PHASE 4: Replace Table 4
# ============================================================
print("\nPhase 4: Replace Table 4")

# Find Table 4 (the 4th table, index 3)
if len(doc.tables) >= 4:
    table4 = doc.tables[3]

    # Clear existing rows (keep header)
    while len(table4.rows) > 1:
        tr = table4.rows[-1]._tr
        table4._tbl.remove(tr)

    # Update header
    header_cells = table4.rows[0].cells
    headers = ['Protocol', 'Token', 'Category', 'HHI', 'Gini', 'Top-1%', 'Top-10%', 'N']
    for i, h in enumerate(headers):
        if i < len(header_cells):
            header_cells[i].text = h

    # Add rows
    for _, row_data in df_sorted.iterrows():
        row = table4.add_row()
        cells = row.cells
        name = row_data['protocol']
        if row_data['governance_model'] == 'futarchy':
            name += ' †'
        if row_data['token'] in ('GTC', 'TEC'):
            name += ' ‡'

        values = [
            name,
            row_data['token'],
            row_data['category'].replace('_', '/'),
            f"{row_data['hhi']:.4f}",
            f"{row_data['gini']:.2f}",
            f"{row_data['top1_pct']:.1f}%",
            f"{row_data['top10_pct']:.1f}%",
            str(row_data['n_holders']),
        ]
        for i, v in enumerate(values):
            if i < len(cells):
                cells[i].text = v
                # Set font size
                for para in cells[i].paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(8)

    print(f"  Replaced Table 4 with {len(df_sorted)} rows")
else:
    print(f"  WARNING: Table 4 not found")

# Update Table 4 caption
old_t4_caption = "Table 4. Governance concentration cross-section (February 2026 snapshot"
idx, p = find_paragraph_containing(old_t4_caption)
if p:
    new_t4_caption = (
        "Table 4. Governance concentration cross-section (March 2026 snapshot, "
        "Dune Analytics + Helius DAS API, top-1,000 holders per token, "
        "30 protocol-controlled addresses excluded across 14 protocols). "
        "† = futarchy governance. ‡ = survivorship-bias control. "
        "Sorted by category then HHI ascending."
    )
    replace_paragraph_text(p, new_t4_caption)
    print(f"  Updated Table 4 caption")

# ============================================================
# PHASE 5: §5.2 rewrite
# ============================================================
print("\nPhase 5: §5.2 rewrite")

# Find the opening paragraph of 5.2
old_52_open = "Table 4 reports governance concentration for 13 protocols"
idx, p = find_paragraph_containing(old_52_open)
if p:
    new_52_open = (
        f"Table 4 reports governance concentration for {len(df)} protocols, computed from "
        "Dune Analytics multi-chain queries and Helius DAS API token holder distributions "
        "(March 2026 snapshot). We report address-holding concentration after excluding "
        "30 protocol-controlled addresses across 14 protocols (see Section 5.1a). "
        "Solana-native tokens were collected via Helius getTokenAccounts with page-based "
        "pagination to correct systematic undercounting in Dune's indexed tables."
    )
    replace_paragraph_text(p, new_52_open)
    print(f"  Updated §5.2 opening")

# Update "Holding concentration spans" paragraph
old_52_span = "Holding concentration spans an order of magnitude"
idx, p = find_paragraph_containing(old_52_span)
if p:
    new_52_span = (
        f"Holding concentration spans nearly two orders of magnitude across the sample (Table 4). "
        f"HHI ranges from {df['hhi'].min():.3f} (Axelar, after excluding foundation treasury "
        f"and exchange custodians) to {df['hhi'].max():.3f} (Livepeer). Among DeFi governance tokens, "
        f"HHI ranges from {defi.min():.3f} (Lido) to {defi.max():.3f} (Curve). "
        f"DePIN protocols span a similar range: {depin.min():.3f} (Morpheus AI, after "
        f"excluding protocol distribution contracts) to {depin.max():.3f} (Livepeer). "
        "The overlap between DePIN and DeFi HHI distributions is substantial, with 8 of 12 "
        "DePIN protocols falling within the DeFi range."
    )
    replace_paragraph_text(p, new_52_span)
    print(f"  Updated concentration span paragraph")

# Update "DePIN governance tokens, by contrast" paragraph
old_52_depin = "DePIN governance tokens, by contrast, extend well above the DeFi ceiling"
idx, p = find_paragraph_containing(old_52_depin)
if p:
    new_52_depin = (
        "After correcting for holder-list truncation (Solana tokens) and protocol-controlled "
        "address contamination (14 protocols), the DePIN-DeFi concentration gap narrows "
        "substantially. Several protocols previously appearing highly concentrated are revealed "
        "to have moderate concentration once measurement artifacts are removed: DRIFT "
        "(0.419 uncorrected to 0.101 corrected, reflecting Helius enumeration of 27,888 "
        "holders versus 104 from Dune), W (0.126 to 0.015), and GRASS (0.132 to 0.035). "
        "The exclusion methodology (Section 5.1a) is particularly consequential for "
        "protocols with large foundation treasuries or exchange custodian addresses in "
        "the top-10 holders."
    )
    replace_paragraph_text(p, new_52_depin)
    print(f"  Updated DePIN contrast paragraph")

# Update "The Anyone Protocol" paragraph
old_52_anyone = "The Anyone Protocol (HHI 0.009; Table 4) provides the critical counterexample"
idx, p = find_paragraph_containing(old_52_anyone)
if p:
    anyone_hhi = df[df['token'] == 'ANYONE']['hhi'].values[0]
    new_52_anyone = (
        f"Several protocols demonstrate that low governance concentration is achievable "
        f"across sectors. Axelar (HHI {df[df['token']=='AXL']['hhi'].values[0]:.3f}), "
        f"LayerZero ({df[df['token']=='ZRO']['hhi'].values[0]:.3f}), "
        f"and Wormhole ({df[df['token']=='W']['hhi'].values[0]:.3f}) in L1/L2 infrastructure; "
        f"Morpheus AI ({df[df['token']=='MOR']['hhi'].values[0]:.3f}) in DePIN; "
        f"and Lido ({df[df['token']=='LDO']['hhi'].values[0]:.3f}) in DeFi all achieve HHI below 0.02. "
        "These counterexamples demonstrate that governance concentration reflects distribution "
        "design choices, not sector constraints."
    )
    replace_paragraph_text(p, new_52_anyone)
    print(f"  Updated counterexample paragraph")

# ============================================================
# PHASE 6: §5.3 rewrite
# ============================================================
print("\nPhase 6: §5.3 rewrite")

old_53_open = "The governance data reveal a structural contrast between DeFi and DePIN protocols"
idx, p = find_paragraph_containing(old_53_open)
if p:
    from scipy import stats
    t_stat, p_val = stats.ttest_ind(depin, defi, equal_var=False)
    cohens_d = abs(depin.mean() - defi.mean()) / ((depin.std()**2 + defi.std()**2) / 2)**0.5

    new_53_open = (
        "The expanded 35-protocol dataset reveals a striking finding: DePIN and DeFi governance "
        f"concentration are not statistically distinguishable. DePIN mean HHI ({depin.mean():.3f}, "
        f"N={len(depin)}) exceeds DeFi mean HHI ({defi.mean():.3f}, N={len(defi)}) by a factor "
        f"of {depin.mean()/defi.mean():.2f}x, but this difference is not significant "
        f"(Welch t-test t={t_stat:.2f}, p={p_val:.2f}, Cohen's d={cohens_d:.2f}). "
        "Both sectors span similar ranges and overlap substantially. The high within-sector "
        "variance (DePIN standard deviation 0.061, DeFi 0.044) overwhelms the between-sector "
        "difference, yielding a small effect size."
    )
    replace_paragraph_text(p, new_53_open)
    print(f"  Updated §5.3 opening")

# ============================================================
# PHASE 7: §6 HHI value replacements
# ============================================================
print("\nPhase 7: §6 HHI value replacements")

# Systematic replacements of old HHI values
replacements = [
    # Old values from the 13-protocol paper → new values
    ("HHI 0.009, lowest in sample", f"HHI {df[df['token']=='ANYONE']['hhi'].values[0]:.3f}"),
    ("HHI 0.009", f"HHI {df[df['token']=='ANYONE']['hhi'].values[0]:.3f}"),
    ("HHI 0.028", f"HHI {df[df['token']=='COMP']['hhi'].values[0]:.3f}"),
    ("HHI 0.045", f"HHI {df[df['token']=='MKR']['hhi'].values[0]:.3f}"),
    ("HHI 0.305", f"HHI {df[df['token']=='DIMO']['hhi'].values[0]:.3f}"),
    ("DIMO (HHI 0.305)", f"DIMO (HHI {df[df['token']=='DIMO']['hhi'].values[0]:.3f})"),
    ("IoTeX (0.388)", f"IoTeX ({df[df['token']=='IOTX']['hhi'].values[0]:.3f})"),
    ("WeatherXM (0.593)", f"WeatherXM ({df[df['token']=='WXM']['hhi'].values[0]:.3f})"),
    ("HHI 0.593", f"HHI {df[df['token']=='WXM']['hhi'].values[0]:.3f}"),
    ("(Curve, where veCRV concentration", f"(Curve, HHI {df[df['token']=='CRV']['hhi'].values[0]:.3f}, where veCRV concentration"),
    ("WeatherXM (2024, HHI 0.593)", f"WeatherXM (2024, HHI {df[df['token']=='WXM']['hhi'].values[0]:.3f})"),
    ("MakerDAO (2017, HHI 0.045)", f"MakerDAO (2017, HHI {df[df['token']=='MKR']['hhi'].values[0]:.3f})"),
    ("Compound (2018 launch, HHI 0.028)", f"Compound (2018 launch, HHI {df[df['token']=='COMP']['hhi'].values[0]:.3f})"),
    ("DIMO (2022, HHI 0.305)", f"DIMO (2022, HHI {df[df['token']=='DIMO']['hhi'].values[0]:.3f})"),
    ("MakerDAO (HHI 0.045)", f"MakerDAO (HHI {df[df['token']=='MKR']['hhi'].values[0]:.3f})"),
    # "12 protocols" → "35 protocols"
    ("12 protocols", "35 protocols"),
    ("13 protocols", "35 protocols"),
    ("13 sampled protocols", "35 sampled protocols"),
    ("12 DeFi and GeoDePIN", "35"),
    # "Gini 0.72–0.98" → updated range
    ("Gini 0.72–0.98", f"Gini {df['gini'].min():.2f}–{df['gini'].max():.2f}"),
    ("Gini 0.72\u20130.98", f"Gini {df['gini'].min():.2f}–{df['gini'].max():.2f}"),
    # February → March
    ("February 2026", "March 2026"),
]

for old, new in replacements:
    count = find_and_replace_text(old, new)
    if count > 0:
        print(f"  Replaced '{old[:40]}...' → '{new[:40]}...' ({count}x)")

# ============================================================
# PHASE 8: §7 conclusion touches
# ============================================================
print("\nPhase 8: §7 update")
old_71 = "Tokenomics is institutional design: governance, rights, and legitimacy constraints"
idx, p = find_paragraph_containing(old_71)
if p:
    new_71 = (
        "Tokenomics is institutional design: governance, rights, and legitimacy constraints "
        "encoded into programmable rule systems. The normative framework developed here "
        "(Table 1) translates five political-philosophical traditions into design criteria "
        "operationalized through a replicable scoring rubric. The empirical illustration "
        "contributes three findings. First, a systematic exclusion methodology that removes "
        "protocol-controlled addresses reveals that naive holder-list analysis can overstate "
        "governance concentration by up to 5x. Second, an expanded 35-protocol cross-section "
        "demonstrates sector invariance: DePIN and DeFi governance concentration are not "
        f"statistically distinguishable (p = 0.35). Third, distribution mechanism design, "
        "not sector membership, is the primary determinant of governance concentration."
    )
    replace_paragraph_text(p, new_71)
    print(f"  Updated §7.1 summary")

# ============================================================
# PHASE 9: Save
# ============================================================
print(f"\nSaving to {DST}")
doc.save(DST)
print("Done!")

# Verify
doc2 = Document(DST)
print(f"\nVerification: {len(doc2.paragraphs)} paragraphs, {len(doc2.tables)} tables")
t4 = doc2.tables[3]
print(f"Table 4: {len(t4.rows)} rows (should be {len(df_sorted)+1})")

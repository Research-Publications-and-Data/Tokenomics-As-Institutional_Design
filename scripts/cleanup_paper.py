"""
B2 Cleanup Pass: Stale Values + Missing Sections
Input:  /Users/zach/B2_Rewrite_34Protocol.docx
Output: /Users/zach/B2_Rewrite_CLEAN.docx

Applies:
  A1: Fix Gini/HHI range in §4.2
  A2: Fix Anyone exclusion sentence in §4.2
  A3: Fix Table 2 (scoring illustration) HHI values for Helium and Anyone
  A4: Fix "HHI above 0.25" paragraph in §6.1 (Para 256)
  A5: Fix WXM/DIMO reference in §6.2 (Para 266)
  A6: Fix CRV HHI in §6.2 (Para 263)
  B:  Insert new §5.4 Voting Power vs Holding Concentration (before old 5.4)
  Rename old 5.4 → 5.5
  C:  Insert Table 6 (Voting vs Holding HHI) after new §5.4 text
  D:  Table 6 uses corrected LDO holding HHI (0.018, actual 0.0185)
"""

import copy
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import lxml.etree as etree

INPUT = '/Users/zach/B2_Rewrite_34Protocol.docx'
OUTPUT = '/Users/zach/B2_Rewrite_CLEAN.docx'


# ── helpers ─────────────────────────────────────────────────────────────────

def replace_in_para(para, old, new):
    """Replace old→new across all runs in a paragraph, preserving first run's format."""
    full = para.text
    if old not in full:
        return False
    new_text = full.replace(old, new)
    # Rewrite into first run, clear the rest
    if para.runs:
        para.runs[0].text = new_text
        for r in para.runs[1:]:
            r.text = ''
    return True


def replace_in_cell(cell, old, new):
    """Replace old→new in all paragraphs within a table cell."""
    changed = False
    for para in cell.paragraphs:
        if replace_in_para(para, old, new):
            changed = True
    return changed


def set_cell_text(cell, text, bold=False, font_size=10):
    """Set cell text with formatting."""
    cell.text = ''
    para = cell.paragraphs[0]
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)


def add_table_after_para(doc, para_index, headers, rows, caption_text):
    """Insert a table + caption after the paragraph at para_index."""
    # We need to insert into the document body XML directly
    body = doc.element.body
    paragraphs = doc.paragraphs

    # Get the XML element of the reference paragraph
    ref_elem = paragraphs[para_index]._element

    # Create caption paragraph
    cap_para = OxmlElement('w:p')
    cap_pPr = OxmlElement('w:pPr')
    cap_pStyle = OxmlElement('w:pStyle')
    cap_pStyle.set(qn('w:val'), 'Caption')
    cap_pPr.append(cap_pStyle)
    cap_para.append(cap_pPr)
    cap_r = OxmlElement('w:r')
    cap_t = OxmlElement('w:t')
    cap_t.text = caption_text
    cap_r.append(cap_t)
    cap_para.append(cap_r)

    # Build table XML
    tbl = OxmlElement('w:tbl')

    # Table properties
    tblPr = OxmlElement('w:tblPr')
    tblStyle = OxmlElement('w:tblStyle')
    tblStyle.set(qn('w:val'), 'TableGrid')
    tblPr.append(tblStyle)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '0')
    tblW.set(qn('w:type'), 'auto')
    tblPr.append(tblW)
    tbl.append(tblPr)

    # Column grid
    tblGrid = OxmlElement('w:tblGrid')
    for _ in headers:
        gridCol = OxmlElement('w:gridCol')
        tblGrid.append(gridCol)
    tbl.append(tblGrid)

    def make_row(cells_text, is_header=False):
        tr = OxmlElement('w:tr')
        for txt in cells_text:
            tc = OxmlElement('w:tc')
            tcPr = OxmlElement('w:tcPr')
            tc.append(tcPr)
            p = OxmlElement('w:p')
            r = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            if is_header:
                b = OxmlElement('w:b')
                rPr.append(b)
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), '20')  # 10pt
            rPr.append(sz)
            r.append(rPr)
            t = OxmlElement('w:t')
            t.text = str(txt)
            r.append(t)
            p.append(r)
            tc.append(p)
            tr.append(tc)
        return tr

    tbl.append(make_row(headers, is_header=True))
    for row in rows:
        tbl.append(make_row(row))

    # Insert after ref_elem: caption then table
    ref_elem.addnext(tbl)
    ref_elem.addnext(cap_para)

    return tbl


def insert_section_before_para(doc, para_index, heading_text, body_paras):
    """Insert heading + body paragraphs before the paragraph at para_index."""
    body_elem = doc.element.body
    ref_elem = doc.paragraphs[para_index]._element

    # We build elements in reverse order and insert each before ref_elem
    # So last body para first, then earlier body paras, then heading
    all_texts = [(heading_text, 'Heading 2')] + [(t, 'Normal') for t in body_paras]

    # Build XML elements
    elems = []
    for text, style_name in all_texts:
        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        pStyle = OxmlElement('w:pStyle')
        pStyle.set(qn('w:val'), style_name.replace(' ', ''))
        pPr.append(pStyle)
        p.append(pPr)
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = text
        r.append(t)
        p.append(r)
        elems.append(p)

    # Insert in order: insert each before ref_elem, so first item ends up first
    # We insert in reverse so final order is correct
    for elem in reversed(elems):
        ref_elem.addnext(elem)
    # Now elems[0] (heading) is right before ref_elem? No — we need them inserted before ref_elem
    # Actually addnext inserts AFTER, so we need addprevious or insert differently
    # Let's use a different approach: insert before
    for elem in elems:
        ref_elem.addprevious(elem)
    # But addprevious of elems in order will put them all before ref_elem in order
    # Actually addprevious inserts the element immediately before ref_elem each time
    # so inserting elem[0], then elem[1], etc. will result in: elem[0], elem[1], ..., ref_elem
    # That's correct!

    return elems


# ── main ────────────────────────────────────────────────────────────────────

def main():
    import shutil
    shutil.copy(INPUT, OUTPUT)
    doc = Document(OUTPUT)

    paras = doc.paragraphs
    print(f"Loaded {len(paras)} paragraphs, {len(doc.tables)} tables")

    # ── A1: Fix Gini/HHI range in §4.2 (Para 183) ──────────────────────────
    old_183 = "inequality (Gini) is near-universal across both DeFi and DePIN (range 0.72 to 0.98), while concentration (HHI) varies by an order of magnitude (0.028 to 0.593)."
    new_183 = "inequality (Gini) is near-universal across all sectors (range 0.45 to 0.99), while concentration (HHI) varies by an order of magnitude (0.004 to 0.199 after excluding protocol-controlled addresses)."
    ok = replace_in_para(paras[183], old_183, new_183)
    print(f"A1 (§4.2 Gini/HHI range): {'OK' if ok else 'NOT FOUND'}")

    # ── A2: Fix Anyone exclusion sentence (Para 184) ────────────────────────
    old_184 = "Results are qualitatively unchanged when excluding Anyone from the GeoDePIN subset (remaining range: HHI 0.119 to 0.593, N=4)."
    new_184 = "Five DePIN protocols (Render 0.032, Grass 0.035, DIMO 0.038, Anyone 0.040, Filecoin 0.047) achieve concentration levels at or below the DeFi median, confirming that low concentration is achievable in hardware-coordination protocols."
    ok = replace_in_para(paras[184], old_184, new_184)
    print(f"A2 (§4.2 Anyone sentence): {'OK' if ok else 'NOT FOUND'}")

    # ── A3: Fix Table 2 (scoring illustration) — Table index 1 ─────────────
    # Helium row: "HHI 0.593" → "HHI 0.102"
    # Anyone row: "Low HHI (0.009)" → "Low HHI (0.040)"
    t1 = doc.tables[1]
    helium_fixed = False
    anyone_fixed = False
    for row in t1.rows:
        for cell in row.cells:
            if 'HHI 0.593' in cell.text:
                replace_in_cell(cell, 'HHI 0.593', 'HHI 0.102')
                helium_fixed = True
            if 'Low HHI (0.009)' in cell.text:
                replace_in_cell(cell, 'Low HHI (0.009)', 'Low HHI (0.040)')
                anyone_fixed = True
    print(f"A3 (Table 2 Helium HHI): {'OK' if helium_fixed else 'NOT FOUND'}")
    print(f"A3 (Table 2 Anyone HHI): {'OK' if anyone_fixed else 'NOT FOUND'}")

    # ── A4: Fix "HHI above 0.25" in §6.1 (Para 256) ────────────────────────
    # Use actual CSV values: LPT=0.199, WXM=0.148, GEOD=0.133
    old_a4 = ("Protocols with HHI above 0.25 (DIMO at 0.305, IoTeX at 0.388, WeatherXM at 0.593) "
              "operate in a regime where Kantian publicity becomes performative rather than substantive: "
              "governance decisions effectively reflect a single entity\u2019s preferences, making "
              "\u201cpublic justification\u201d a formality.")
    new_a4 = ("Protocols at the upper end of the concentration range (Livepeer at 0.199, WeatherXM at 0.148, "
              "GEODNET at 0.133) exhibit concentration levels the normative framework flags for scrutiny, "
              "though no protocol in the corrected cross-section exceeds the conventional 0.25 antitrust "
              "threshold. Even at these levels, governance decisions can disproportionately reflect a small "
              "set of holders\u2019 preferences, making \u201cpublic justification\u201d structurally thin.")
    ok = replace_in_para(paras[256], old_a4, new_a4)
    if not ok:
        # Try simpler fragment
        ok = replace_in_para(paras[256],
            "Protocols with HHI above 0.25 (DIMO at 0.305, IoTeX at 0.388, WeatherXM at 0.593) operate in a regime where Kantian publicity becomes performative rather than substantive",
            "Protocols at the upper end of the concentration range (Livepeer at 0.199, WeatherXM at 0.148, GEODNET at 0.133) exhibit concentration levels the normative framework flags for scrutiny, though no protocol in the corrected cross-section exceeds the conventional 0.25 antitrust threshold")
    print(f"A4 (§6.1 HHI above 0.25): {'OK' if ok else 'NOT FOUND'}")

    # ── A5: Fix WXM/DIMO in §6.2 (Para 266) ────────────────────────────────
    # Actual CSV: WXM=0.148 (keep), DIMO=0.038 (fix from 0.305)
    old_a5 = "WXM holding HHI 0.148; DIMO 0.305"
    new_a5 = "WXM holding HHI 0.148; DIMO 0.038"
    ok = replace_in_para(paras[266], old_a5, new_a5)
    print(f"A5 (§6.2 WXM/DIMO HHI): {'OK' if ok else 'NOT FOUND'}")

    # ── A6: Fix CRV HHI in §6.2 (Para 263) ─────────────────────────────────
    # Actual CSV: CRV HHI = 0.1706 → 0.171
    old_a6 = "CRV (HHI 0.174)"
    new_a6 = "CRV (HHI 0.171)"
    ok = replace_in_para(paras[263], old_a6, new_a6)
    print(f"A6 (§6.2 CRV HHI): {'OK' if ok else 'NOT FOUND'}")

    # ── Renumber old 5.4 → 5.5 (Para 248) ──────────────────────────────────
    old_54 = "5.4 Hypothesis Evidence Status"
    new_55 = "5.5 Hypothesis Evidence Status"
    ok = replace_in_para(paras[248], old_54, new_55)
    print(f"Renumber §5.4→§5.5: {'OK' if ok else 'NOT FOUND'}")

    # ── B: Insert new §5.4 Voting Power vs Holding Concentration ────────────
    # Insert before para 248 (the old 5.4, now 5.5)
    # Para 248 is "5.4 Hypothesis Evidence Status" (now renamed above to 5.5)

    sec54_heading = "5.4 Voting Power vs Holding Concentration"

    sec54_body = [
        (
            "The holding-based HHI reported in Table 4 measures token distribution, not governance power. "
            "For protocols with delegation, a small number of delegates may control voting power far exceeding "
            "their personal holdings. To quantify this gap, we collected delegated voting power from Tally "
            "(on-chain governance: Compound, Aave, Uniswap, Optimism, Arbitrum) and Snapshot (off-chain "
            "governance: DIMO, WeatherXM, Lido) for the 8 protocols with sufficient governance activity in "
            "the 12 months preceding the snapshot."
        ),
        (
            "Table 6 reports the holding-voting HHI gap. The delegation amplification ratio (voting HHI "
            "divided by holding HHI) reveals a sector-dependent pattern absent from the holding data alone. "
            "DePIN protocols show extreme delegation amplification: DIMO\u2019s voting HHI (0.228) is 6.0 times "
            "its holding HHI (0.038), and WeatherXM\u2019s (0.486) is 3.3 times its holding HHI (0.148). A small "
            "number of delegates, likely foundation-affiliated or early-investor wallets, control governance "
            "outcomes despite relatively distributed token holdings."
        ),
        (
            "DeFi protocols show a mixed pattern. Lido exhibits substantial amplification (4.8 times), "
            "consistent with its known delegate concentration. Compound shows moderate amplification "
            "(1.9 times). Aave is roughly neutral (1.3 times). Uniswap shows mild distribution through "
            "delegation (0.67 times), suggesting its active delegate ecosystem partially decompresses "
            "governance power."
        ),
        (
            "Infrastructure protocols show the reverse pattern: delegation distributes rather than "
            "concentrates. Optimism\u2019s voting HHI (0.033) is 0.27 times its holding HHI (0.121), and "
            "Arbitrum\u2019s (0.052) is 0.54 times its holding HHI (0.097). Both protocols have invested "
            "heavily in delegate programs, compensated delegation, and constitutional governance structures "
            "that effectively distribute voting power beyond holding concentration."
        ),
        (
            "The delegation amplification finding reintroduces a sector-dependent pattern that the holding "
            "data alone did not support. While holding concentration is sector-invariant (Section 5.3), "
            "delegation amplification is not: DePIN protocols concentrate governance power through delegation "
            "at rates 3 to 6 times higher than their holding concentration, while infrastructure protocols "
            "achieve the opposite. This pattern is consistent with differences in governance infrastructure "
            "maturity."
        ),
    ]

    # Find para 248 index in current doc (it was just renamed to 5.5)
    # Use the renamed text to locate it
    para_248_idx = None
    for i, p in enumerate(doc.paragraphs):
        if "5.5 Hypothesis Evidence Status" in p.text:
            para_248_idx = i
            break
    if para_248_idx is None:
        print("ERROR: could not find renamed 5.5 paragraph")
    else:
        print(f"Found §5.5 at para index {para_248_idx}")

    # Insert new §5.4 before that paragraph
    # Use addprevious on the reference element
    ref_elem = doc.paragraphs[para_248_idx]._element

    def make_para_elem(text, style='Normal'):
        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        pStyle = OxmlElement('w:pStyle')
        pStyle.set(qn('w:val'), style.replace(' ', ''))
        pPr.append(pStyle)
        p.append(pPr)
        r = OxmlElement('w:r')
        t_elem = OxmlElement('w:t')
        t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t_elem.text = text
        r.append(t_elem)
        p.append(r)
        return p

    def make_table_elem(headers, rows):
        tbl = OxmlElement('w:tbl')
        tblPr = OxmlElement('w:tblPr')
        tblStyle = OxmlElement('w:tblStyle')
        tblStyle.set(qn('w:val'), 'TableGrid')
        tblPr.append(tblStyle)
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:w'), '0')
        tblW.set(qn('w:type'), 'auto')
        tblPr.append(tblW)
        tbl.append(tblPr)
        tblGrid = OxmlElement('w:tblGrid')
        for _ in headers:
            gridCol = OxmlElement('w:gridCol')
            tblGrid.append(gridCol)
        tbl.append(tblGrid)

        def make_row_elem(cells_text, is_header=False):
            tr = OxmlElement('w:tr')
            for txt in cells_text:
                tc = OxmlElement('w:tc')
                tcPr = OxmlElement('w:tcPr')
                tc.append(tcPr)
                p = OxmlElement('w:p')
                r = OxmlElement('w:r')
                rPr = OxmlElement('w:rPr')
                if is_header:
                    b = OxmlElement('w:b')
                    rPr.append(b)
                sz = OxmlElement('w:sz')
                sz.set(qn('w:val'), '20')
                rPr.append(sz)
                r.append(rPr)
                t = OxmlElement('w:t')
                t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                t.text = str(txt)
                r.append(t)
                p.append(r)
                tc.append(p)
                tr.append(tc)
            return tr

        tbl.append(make_row_elem(headers, is_header=True))
        for row in rows:
            tbl.append(make_row_elem(row))
        return tbl

    # Build Table 6 data (D: corrected holding HHI values)
    # LDO actual: 0.018498 → display as 0.018; ratio = 0.088/0.018498 = 4.76 ≈ 4.8x
    table6_headers = ['Protocol', 'Category', 'Holding HHI', 'Voting HHI', 'Ratio', 'Source']
    table6_rows = [
        ['DIMO',       'DePIN', '0.038', '0.228', '6.0x', 'Snapshot'],
        ['Lido',       'DeFi',  '0.018', '0.088', '4.8x', 'Snapshot'],  # corrected from 0.013
        ['WeatherXM',  'DePIN', '0.148', '0.486', '3.3x', 'Snapshot'],
        ['Compound',   'DeFi',  '0.028', '0.053', '1.9x', 'Tally'],
        ['Aave',       'DeFi',  '0.059', '0.076', '1.3x', 'Tally'],
        ['Uniswap',    'DeFi',  '0.101', '0.068', '0.67x', 'Tally'],
        ['Arbitrum',   'Infra', '0.097', '0.052', '0.54x', 'Tally'],
        ['Optimism',   'Infra', '0.121', '0.033', '0.27x', 'Tally'],
    ]

    table6_caption = (
        "Table 6. Holding versus voting concentration for protocols with sufficient governance activity "
        "(12-month lookback). Voting HHI computed from Tally delegate voting power (on-chain) or "
        "Snapshot vote-weighted power (off-chain). Ratio above 1.0 indicates delegation concentrates "
        "governance; below 1.0 indicates delegation distributes."
    )

    # Insert in this order before ref_elem (using addprevious, which inserts immediately before):
    # heading, body para 1, body para 2, ..., body para 5, table caption, table 6
    # Since addprevious(x) puts x immediately before ref_elem, we insert in order:
    elements_to_insert = []
    elements_to_insert.append(make_para_elem(sec54_heading, 'Heading2'))
    for body_text in sec54_body:
        elements_to_insert.append(make_para_elem(body_text, 'Normal'))
    elements_to_insert.append(make_para_elem(table6_caption, 'Normal'))
    elements_to_insert.append(make_table_elem(table6_headers, table6_rows))

    for elem in elements_to_insert:
        ref_elem.addprevious(elem)

    print(f"B+C: Inserted §5.4 ({len(sec54_body)} body paras) + Table 6 ({len(table6_rows)} rows)")

    # ── Save ────────────────────────────────────────────────────────────────
    doc.save(OUTPUT)
    print(f"\nSaved to {OUTPUT}")
    print(f"Final: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")

    # Quick verification
    content = ' '.join(p.text for p in doc.paragraphs)
    checks = [
        ("A1 Gini fix",     "0.45 to 0.99" in content),
        ("A1 HHI fix",      "0.004 to 0.199" in content),
        ("A1 old gone",     "0.72 to 0.98" not in content),
        ("A2 new sentence", "Five DePIN protocols" in content),
        ("A2 old gone",     "HHI 0.119 to 0.593" not in content),
        ("A3 Helium HHI",   "HHI 0.102" in content),
        ("A3 old Helium",   "HHI 0.593" not in content),
        ("A3 Anyone HHI",   "Low HHI (0.040)" in content),
        ("A4 new §6.1",     "Livepeer at 0.199" in content),
        ("A4 old gone",     "DIMO at 0.305" not in content),
        ("A5 DIMO fix",     "DIMO 0.038" in content),
        ("A5 old gone",     "DIMO 0.305" not in content),
        ("A6 CRV fix",      "CRV (HHI 0.171)" in content),
        ("B §5.4 heading",  "5.4 Voting Power" in content),
        ("B §5.5 renumber", "5.5 Hypothesis Evidence" in content),
    ]
    # Also check tables
    for t in doc.tables:
        for row in t.rows:
            cells = [c.text for c in row.cells]
            if 'Lido' in cells:
                lido_holding = cells[cells.index('Lido') + 1] if cells.index('Lido') + 1 < len(cells) else ''
                # check via row text
                row_text = ' '.join(cells)
                if '0.018' in row_text and 'Lido' in row_text:
                    checks.append(("D Lido HHI corrected", True))
                    break

    print("\n=== VERIFICATION ===")
    all_pass = True
    for name, result in checks:
        status = "PASS" if result else "FAIL"
        print(f"  [{status}] {name}")
        if not result:
            all_pass = False
    print(f"\n{'ALL CHECKS PASSED' if all_pass else 'SOME CHECKS FAILED'}")


if __name__ == '__main__':
    main()
